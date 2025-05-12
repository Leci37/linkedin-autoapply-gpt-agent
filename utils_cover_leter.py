import os
from docx import Document
from selenium.webdriver.common.by import By

from utils_move import gentle_scroll
from GPT_Assistant_utils import Get_answerGPT_for_question

def generate_cover_letter_text(job_data, assistant):
    """
    Asks GPT to return a concise, formatted cover letter using job details, ending with Luis Lecinana's contact info.
    Output uses markdown formatting (**bold**, *italic*) and ends with a fixed signature block.
    """
    job_title = job_data['job']['title']
    company_name = job_data['company']['name']
    city = job_data['job']['cities_available'][0] if job_data['job']['cities_available'] else "España"
    description = job_data['job']['full_description'][:1000]
    lang = job_data['job'].get('description_language', 'English')

    lang_prompt = {
        "Spanish": "en español",
        "English": "in English",
        # "French": "en français",
        # "German": "auf Deutsch"
    }.get(lang, "in English")

    contact_footer = (
        "\n\nGracias por su tiempo.\n"
        "Atentamente,\n"
        "Luis Lecinana\n"
        "LinkedIn: http://linkedin.com/in/xxxx\n"
        "GitHub: http://github.com/Leci37\n"
        "Email: x.xxxx@gmail.com\n"
        "Teléfono: +34 xxx xxx xxx"
    ) if lang == "Spanish" else (
        "\n\nThank you very much for your time.\n"
        "Sincerely,\n"
        "Luis Lecinana\n"
        "LinkedIn: http://linkedin.com/in/xxxx\n"
        "GitHub: http://github.com/Leci37\n"
        "Email: x.xxxx@gmail.com\n"
        "Phone: +34 xxx xxx xxx"
    )

    prompt = (
        f"Write a Concise and conversational cover letter of about 660–760 characters {lang_prompt} "
        f"for the job '{job_title}' at '{company_name}' in {city}. "
        f"The tone should reflect Luis Lecinana's voice and reference key points from the job description:\n\n"
        f"{description}\n\n"
        f"Use markdown formatting for emphasis (**bold**, *italic*). Return only the body text , add some \\n  "
        f"\n{contact_footer}"
    )

    body = assistant.ask(prompt).strip()
    return body




from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os
import re
import unicodedata

def sanitize_filename(text):
    text = unicodedata.normalize('NFKD', text).encode('ASCII', 'ignore').decode()
    text = re.sub(r"[^\w\s-]", "", text).strip().replace(" ", "_")
    return text.lower()

def save_cover_letter_to_doc(text, job_id, company_name, folder="Data/cover_letters"):
    """
    Saves a markdown-formatted, well-spaced cover letter to .docx.
    """
    os.makedirs(folder, exist_ok=True)
    filename = f"cover_letter_{sanitize_filename(company_name)}_{job_id}.docx"
    path = os.path.join(folder, filename)

    doc = Document()

    # Set style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(10)
    style.paragraph_format.space_before = Pt(6)

    # Split into paragraphs by double newlines
    blocks = re.split(r"\n\s*\n", text.strip())

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(10)
        paragraph.paragraph_format.space_before = Pt(6)

        cursor = 0
        for match in re.finditer(r"(\*\*.+?\*\*|\*.+?\*)", block):
            start, end = match.span()
            match_text = match.group()
            clean_text = match_text.strip("*")

            if start > cursor:
                paragraph.add_run(block[cursor:start])

            run = paragraph.add_run(clean_text)
            if match_text.startswith("**"):
                run.bold = True
            elif match_text.startswith("*"):
                run.italic = True

            cursor = end

        if cursor < len(block):
            paragraph.add_run(block[cursor:])

    doc.save(path)
    return path

import os
import shutil
from selenium.webdriver.common.by import By
from utils_move import gentle_scroll

def upload_cover_letter(driver, original_path):
    """
    Copies the original cover letter and renames it to 'Luis_Lecinana_letter_{rest}.docx',
    then uploads that file via input[type=file] without opening a file dialog.
    """
    try:
        dir_path = os.path.dirname(original_path)
        base_name = os.path.basename(original_path)

        # Extract the suffix from the original file (e.g. 'devoteam_4194879978.docx')
        match = base_name.replace("cover_letter_", "")
        renamed_name = f"Luis_Lecinana_letter_{match}"
        renamed_path = os.path.join(dir_path, renamed_name)

        # Copy and rename the file
        shutil.copy(original_path, renamed_path)

        # Locate file inputs and upload
        inputs = driver.find_elements(By.XPATH, "//input[@type='file']")

        for input_el in inputs:
            if not input_el.is_enabled():
                continue

            driver.execute_script("arguments[0].style.display = 'block';", input_el)
            driver.execute_script("arguments[0].style.visibility = 'visible';", input_el)
            driver.execute_script("arguments[0].style.opacity = 1;", input_el)

            gentle_scroll(driver, input_el)
            input_el.send_keys(os.path.abspath(renamed_path))
            print(f"📤 Uploaded renamed cover letter: {renamed_path}")
            return True

        print("⚠️ No usable file input found.")
        return False

    except Exception as e:
        print(f"❌ Failed to upload cover letter: {e}")
        return False

